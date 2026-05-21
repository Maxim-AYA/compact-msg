"""
Формирование Word-отчёта по жизненному циклу проекта МСГ_RBI Репино Санаторий
из C:\\Авраменко\\Claude Code Projects\\МСГ\\2. Репино\\_lifecycle.json
"""
import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:\Авраменко\Claude Code Projects\МСГ\2. Репино")
DATA = ROOT / "_lifecycle.json"
OUT  = ROOT / "МСГ Репино — Отчёт по жизненному циклу проекта.docx"

GREEN = RGBColor(0x1F, 0x7A, 0x1F)
ORANGE = RGBColor(0xC8, 0x6B, 0x00)
RED = RGBColor(0xB0, 0x10, 0x10)
GREY = RGBColor(0x66, 0x66, 0x66)

SECTION_ORDER = [
    "ПОДГОТОВИТЕЛЬНЫЕ РАБОТЫ", "КОТЛОВАН", "ШПУНТ", "БАШЕННЫЙ КРАН",
    "ФУНДАМЕНТНАЯ ПЛИТА", "ГИДРОИЗОЛЯЦИЯ НИЖЕ ОТМ. 0.000",
    "МОНОЛИТНЫЕ Ж/Б СТЕНЫ НИЖЕ ОТМ. 0.000", "МОНОЛИТНЫЕ Ж/Б ПЕРЕКРЫТИЯ",
    "МОНОЛИТНЫЕ Ж/Б СТЕНЫ ВЫШЕ ОТМ. 0.000", "Ж/Б ЛЕСТНИЦЫ",
    "Монтаж сборных элементов каркаса", "Устройство внутренних перегородок",
    "Устройство кровли здания", "Монтаж оконных и балконных блоков из алюминиевого профиля",
    "Монтаж витражных конструкций", "Устройство фасадов и козырьков",
    "Монтаж лифтового оборудования", "Монтаж дверных блоков",
    "Монтаж металлических конструкций", "Устройство стяжек полов",
    "Устройство систем отопления", "Монтаж ИТП",
    "Устройство систем водоотведения", "Устройство систем водоснабжения",
    "Устройство систем вентиляции", "Внутренние отделочные работы",
    "Внутренние электромонтажные работы (ЭС)", "Внутренние слаботочные работы (слаботочка)",
    "Наружные инженерные сети - выполняются силами Заказчика",
    "Благоустройство", "ВВОД ОБЪЕКТА В ЭКСПЛУАТАЦИЮ",
]

def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "808080")
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def add_table(doc, headers, rows, col_widths=None, header_fill="1F3A5F"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = "Calibri"
        shade_cell(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            text = val if isinstance(val, str) else str(val)
            color = None
            if isinstance(val, tuple):  # (text, color)
                text, color = val
            r = p.add_run(text); r.font.size = Pt(10); r.font.name = "Calibri"
            if color: r.font.color.rgb = color
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl

def fmt_pct(v): return f"{v:.1f}%"

def status_color(p):
    if p is None: return GREY
    if p >= 99.5: return GREEN
    if p <= 0.5:  return RED
    return ORANGE

def make_progress_bar(pct, width=20):
    pct = max(0.0, min(100.0, pct))
    filled = int(round(pct / 100.0 * width))
    return "█" * filled + "░" * (width - filled)

def main():
    data = json.loads(DATA.read_text(encoding="utf-8"))
    doc = Document()

    # --- стили базовые ---
    section = doc.sections[0]
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8);  section.bottom_margin = Cm(1.8)
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

    # --- титул ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ОТЧЁТ ПО ЖИЗНЕННОМУ ЦИКЛУ ПРОЕКТА")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    r.font.name = "Calibri"
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("МСГ_RBI Репино Санаторий")
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    add_para(doc, f"Дата отчёта: {date.today().strftime('%d.%m.%Y')}",
             italic=True, color=GREY, align="center")
    add_para(doc, "Источник: МСГ_RBI Репино Санаторий.xlsx (лист «МСГ, ГПР» + лист «РС»)",
             italic=True, size=9, color=GREY, align="center")

    doc.add_paragraph()

    # --- 1. Резюме ---
    add_heading(doc, "1. Краткое резюме", 1)
    lc = data["lifecycle"]
    smr = next(x for x in lc if x["stage"].startswith("СМР"))
    rd  = next(x for x in lc if x["stage"].startswith("РД"))
    pk  = next(x for x in lc if x["stage"].startswith("Пакеты"))
    tn  = next(x for x in lc if x["stage"].startswith("Тендеры"))
    dg  = next(x for x in lc if x["stage"].startswith("Договоры"))
    fn  = next(x for x in lc if x["stage"].startswith("Финансирование"))
    mb  = next(x for x in lc if x["stage"].startswith("Мобилизация"))

    summary = (
        f"На отчётную дату проект «Репино Санаторий» (6 корпусов + общеплощадочные работы) "
        f"находится на стыке подготовительной фазы (закупки/контрактация) и активной СМР-фазы. "
        f"Рабочая документация выпущена полностью ({rd['done']}/{rd['total']}, {fmt_pct(rd['pct_done'])}), "
        f"пакеты документов — {pk['done']}/{pk['total']} ({fmt_pct(pk['pct_done'])}). "
        f"Контрактация: тендеры закрыты на {fmt_pct(tn['pct_done'])} ({tn['done']} из {tn['total']}), "
        f"договоры заключены на {fmt_pct(dg['pct_done'])} ({dg['done']} из {dg['total']}). "
        f"Мобилизация подрядчиков — {fmt_pct(mb['pct_done'])} ({mb['done']} из {mb['total']}). "
        f"СМР выполнено {fmt_pct(smr['pct_done'])} ({smr['done']} работ из {smr['total']}), "
        f"в производстве — {smr['in_progress']} работ. "
        f"Общая сметная стоимость по 31 лоту — {data['total_cost']/1e9:.2f} млрд руб."
    )
    add_para(doc, summary, size=11)

    # ⚠ ключевые сигналы
    add_para(doc, "Ключевые наблюдения:", bold=True)
    bullets = [
        f"Финансирование (аванс): {fn['done']}/{fn['total']} — статусы оплат в файле не "
        f"проставлены, требуется сверка с финблоком (Гончарова А.).",
        f"Договоры опережают мобилизацию: {dg['done']} заключённых vs. {mb['done']} мобилизаций — "
        f"между этими этапами должно стоять финансирование, его нулевой статус подтверждает разрыв.",
        f"Тендеры в работе: {tn['in_progress']} активных конкурсов — это основная зона работы "
        f"закупочного блока (Лисица Д.) на ближайший период.",
        f"СМР в производстве: {smr['in_progress']} работ — все на этапе закрытия подземной части и старта надземной.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(b); r.font.size = Pt(11); r.font.name = "Calibri"

    # --- 2. Сводка по жизненному циклу ---
    add_heading(doc, "2. Сводка по этапам жизненного цикла", 1)
    add_para(doc, "Каждая работа в графике проходит через 7 этапов: РД → Пакет → Тендер → "
                  "Договор → Финансирование → Мобилизация → СМР. Один тендер/договор может "
                  "закрывать несколько работ, поэтому количество позиций по этапам различается.", size=10)

    rows = []
    for s in lc:
        prog = make_progress_bar(s["pct_done"])
        rows.append([
            s["stage"],
            s["responsible"],
            str(s["total"]),
            (str(s["done"]), GREEN),
            (str(s["in_progress"]), ORANGE),
            (str(s["not_started"]), RED),
            (f"{prog} {fmt_pct(s['pct_done'])}", status_color(s["pct_done"])),
        ])
    add_table(doc, ["Этап", "Ответственный", "Всего", "Заверш.", "В работе", "Не начато", "Готовность"],
              rows, col_widths=[4.2, 2.6, 1.2, 1.4, 1.4, 1.5, 5.0])

    # --- 3. РД и Пакеты ---
    add_heading(doc, "3. Рабочая документация и пакеты", 1)
    add_para(doc,
        f"Рабочая документация выпущена в полном объёме — все {rd['total']} комплектов ({fmt_pct(rd['pct_done'])}) "
        f"переданы строительному блоку. Это ключевой задел, обеспечивающий бесперебойную "
        f"передачу работ в тендерное и контрактное производство.")
    add_para(doc,
        f"Пакеты документов сформированы по {pk['done']} из {pk['total']} позиций "
        f"({fmt_pct(pk['pct_done'])}). Открытыми остаются {pk['not_started']} пакетов — это работы "
        f"конца цикла (отделка, благоустройство, ввод).", size=11)

    # --- 4. Тендеры ---
    add_heading(doc, "4. Тендеры", 1)
    add_para(doc,
        f"Всего {tn['total']} тендерных процедур. Закрыто — {tn['done']} ({fmt_pct(tn['pct_done'])}), "
        f"в активной фазе — {tn['in_progress']}, не начато — {tn['not_started']}. "
        f"Текущее «горлышко» закупочного процесса.", size=11)

    # --- 5. Договоры и финансирование ---
    add_heading(doc, "5. Договоры и финансирование", 1)
    add_para(doc,
        f"Договоры: заключено {dg['done']} из {dg['total']} ({fmt_pct(dg['pct_done'])}); "
        f"в работе — {dg['in_progress']}, не начато — {dg['not_started']}.")
    p = doc.add_paragraph()
    r = p.add_run(
        f"Финансирование (авансы): по факту в системе ноль выполненных платежей "
        f"({fn['done']}/{fn['total']}). При наличии {dg['done']} заключённых договоров такой "
        f"статус выглядит аномально — необходимо синхронизироваться с финблоком "
        f"(Гончарова А.) и уточнить, отражаются ли проведённые авансы в МСГ."
    )
    r.font.size = Pt(11); r.bold = True; r.font.color.rgb = RED; r.font.name = "Calibri"

    # --- 6. Мобилизация ---
    add_heading(doc, "6. Мобилизация подрядчиков и МТР", 1)
    add_para(doc,
        f"Мобилизация выполнена по {mb['done']} позициям из {mb['total']} ({fmt_pct(mb['pct_done'])}). "
        f"Это соответствует фактическому количеству активных подрядчиков на площадке "
        f"и обеспечивает текущий темп СМР по подземной части и старту монолита надземной части.", size=11)

    # --- 7. СМР: разрез по разделам ---
    add_heading(doc, "7. СМР по разделам", 1)
    sections = data["sections_smr"]
    rows = []
    ordered = [s for s in SECTION_ORDER if s in sections]
    for sec in ordered:
        v = sections[sec]
        pct = (100.0 * v["done"] / v["total"]) if v["total"] else 0.0
        rows.append([
            sec,
            str(v["total"]),
            (str(v["done"]), GREEN),
            (str(v["in_progress"]), ORANGE),
            (str(v["not_started"]), RED),
            (f"{make_progress_bar(pct, 12)} {fmt_pct(pct)}", status_color(pct)),
        ])
    # сводная строка
    tot = {"total": 0, "done": 0, "in_progress": 0, "not_started": 0}
    for sec in ordered:
        for k in tot: tot[k] += sections[sec][k]
    pct_all = (100.0 * tot["done"] / tot["total"]) if tot["total"] else 0.0
    rows.append([
        ("ИТОГО", RGBColor(0,0,0)),
        (str(tot["total"]), RGBColor(0,0,0)),
        (str(tot["done"]), GREEN),
        (str(tot["in_progress"]), ORANGE),
        (str(tot["not_started"]), RED),
        (f"{make_progress_bar(pct_all, 12)} {fmt_pct(pct_all)}", status_color(pct_all)),
    ])
    add_table(doc, ["Раздел", "Всего", "Заверш.", "В работе", "Не начато", "Готовность"],
              rows, col_widths=[7.5, 1.3, 1.5, 1.5, 1.5, 4.5])

    # --- 8. СМР: разрез по корпусам ---
    add_heading(doc, "8. СМР по корпусам", 1)
    bld = data["by_building"]
    rows = []
    bld_order = ["К1","К2","К3","К4","К5","К6","БК","ИК","К1-К2","К3-К6","К1-К6","(общее)"]
    for b in bld_order:
        if b not in bld: continue
        v = bld[b]
        pct = (100.0 * v["done"] / v["total"]) if v["total"] else 0.0
        rows.append([
            b, str(v["total"]),
            (str(v["done"]), GREEN),
            (str(v["in_progress"]), ORANGE),
            (str(v["not_started"]), RED),
            (f"{make_progress_bar(pct, 14)} {fmt_pct(pct)}", status_color(pct)),
        ])
    # остатки
    for b, v in bld.items():
        if b in bld_order: continue
        pct = (100.0 * v["done"] / v["total"]) if v["total"] else 0.0
        rows.append([
            b, str(v["total"]),
            (str(v["done"]), GREEN),
            (str(v["in_progress"]), ORANGE),
            (str(v["not_started"]), RED),
            (f"{make_progress_bar(pct, 14)} {fmt_pct(pct)}", status_color(pct)),
        ])
    add_table(doc, ["Корпус", "Всего", "Заверш.", "В работе", "Не начато", "Готовность"],
              rows, col_widths=[2.0, 1.5, 1.5, 1.5, 1.5, 5.5])

    # --- 9. Активные СМР работы ---
    add_heading(doc, "9. Активные работы СМР (в производстве)", 1)
    if data["active_smr"]:
        rows = []
        for w in data["active_smr"]:
            rows.append([
                (w["section"] or "")[:40],
                (w["name"] or "")[:55],
                w["building"] or "",
                f"{w['plan_start']} → {w['plan_end']}" if w["plan_start"] else "",
                (f"{w['fact_pct']:.0f}%", status_color(w["fact_pct"])),
            ])
        add_table(doc, ["Раздел", "Работа", "Корпус", "Срок (план)", "Факт %"],
                  rows, col_widths=[3.5, 6.0, 1.3, 3.0, 1.5])
    else:
        add_para(doc, "Активных работ не зафиксировано.", italic=True, color=GREY)

    # --- 10. Сметная стоимость по лотам ---
    add_heading(doc, "10. Сметная стоимость по лотам (лист «РС»)", 1)
    add_para(doc,
        f"В смете {len(data['lots'])} лотов на {data['total_positions']} позиций. "
        f"Общая стоимость с НДС — {data['total_cost']:,.0f} руб "
        f"({data['total_cost']/1e9:.2f} млрд руб).".replace(",", " "),
        size=11)
    rows = []
    for lot in data["lots"]:
        share = (100.0 * lot["cost_total"] / data["total_cost"]) if data["total_cost"] else 0
        rows.append([
            lot["lot"][:80],
            str(lot["positions"]),
            f"{lot['cost_total']:,.0f}".replace(",", " "),
            f"{share:.1f}%",
        ])
    add_table(doc, ["Лот", "Позиций", "Стоимость, руб", "Доля"],
              rows, col_widths=[10.5, 1.6, 3.5, 1.4])

    # --- Заключение ---
    add_heading(doc, "11. Выводы и фокус ближайшего периода", 1)
    out_text = (
        "1. Документационный задел готов: РД и пакеты — два этапа, не блокирующие проект.\n"
        "2. Узкое место — финансирование. Нулевой статус по авансам в МСГ при ~71 заключённом "
        "договоре требует сверки и уточнения порядка отражения оплат.\n"
        "3. Тендерный блок (Лисица Д.) держит 91 активный конкурс — основная нагрузка "
        "закупки на ближайшие недели.\n"
        "4. По СМР проект на ~10% выполнения; критическая зона — К1, К2 (отставание по "
        "фундаменту и подземной части), К3–К6 ушли вперёд на стены/перекрытия 1–2 этажа.\n"
        "5. Контрольные точки текущего месяца — закрытие фундамента К1/К2, продолжение "
        "монолита надземной части К3–К6, завершение гидроизоляции и обратной засыпки.\n"
    )
    for line in out_text.strip().split("\n"):
        add_para(doc, line, size=11)

    doc.save(OUT)
    print(f"Документ сохранён: {OUT}")
    print(f"Размер: {OUT.stat().st_size:,} байт")

if __name__ == "__main__":
    main()
